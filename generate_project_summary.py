"""
生成《中州养老管理系统 - 项目总结》优化版DOCX
基于原始文档重新整理，补充测试视角的内容
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# 样式设置
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h


def add_bold_text(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    return run


# ========== 封面 ==========
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('中州养老管理系统')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('项 目 总 结')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0, 102, 153)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('项目角色：软件测试工程师\n项目周期：2024.02 - 2024.07')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(102, 102, 102)

doc.add_page_break()

# ========== 正文 ==========

add_heading_styled('一、项目概述', level=1)

doc.add_paragraph(
    '中州养老管理系统是专门为养老机构定制开发的一套专业养老服务管理平台，系统分为两部分：后台管理系统和微信小程序。'
)
doc.add_paragraph(
    '后台管理系统：供养老院员工使用，主要包含老人入住管理、退住管理、护理管理、床位管理、智能监测等核心模块，'
    '帮助养老院员工高效完成日常工作。'
)
doc.add_paragraph(
    '微信小程序：供入住老人及家属使用，功能包括微信登录、预约参观、查看老人健康数据、在线缴费等。'
)

doc.add_paragraph('技术架构：')
techs = [
    ['后端', 'Spring Boot 3.5.x, Spring Security, MyBatis, MyBatis-Plus'],
    ['数据库', 'MySQL 8.x + Redis 7.x（缓存）'],
    ['前端', 'Vue 3 + Element Plus + Axios'],
    ['接口文档', 'SpringDoc OpenAPI 2.x + Apifox'],
    ['AI集成', '百度千帆大模型（健康体检报告智能分析）'],
    ['IoT', '设备数据采集 + RabbitMQ + 线程池异步处理'],
    ['构建工具', 'Maven + Git'],
]

table = doc.add_table(rows=1 + len(techs), cols=2)
table.style = 'Light Grid Accent 1'
table.cell(0, 0).text = '技术维度'
table.cell(0, 1).text = '技术选型'
for i, (k, v) in enumerate(techs):
    table.rows[i + 1].cells[0].text = k
    table.rows[i + 1].cells[1].text = v

doc.add_paragraph()

# ========== 二、项目回顾 ==========
add_heading_styled('二、项目回顾', level=1)

add_heading_styled('2.1 若依框架的使用', level=2)
doc.add_paragraph('项目基于若依框架（RuoYi-Vue）二次开发，充分利用了框架提供的以下核心能力：')
doc.add_paragraph('代码生成功能：利用框架提供的代码生成器，基于数据库表结构自动生成前后端CRUD代码及系统菜单。')
doc.add_paragraph('表单构建功能：通过图形化界面拖拽方式快速生成复杂页面。')
doc.add_paragraph('定时任务功能：在线管理定时任务，支持添加、修改、删除及执行日志查看。')
doc.add_paragraph('权限管理：基于RBAC模型，通过五张核心表（用户、角色、菜单、用户角色关系、角色菜单关系）实现精细化权限控制。')

add_heading_styled('2.2 入住管理模块', level=2)
doc.add_paragraph('入住管理模块是项目的核心模块之一，涉及体检报告、入住登记等功能：')
doc.add_paragraph('体检报告：老人入住养老院之前，需要提供一份体检报告（PDF版本），系统上传体检报告后AI模型会自动分析体检报告并生成健康建议。')
doc.add_paragraph('入住登记：确认老人要入住养老院后，可在系统上发起入住登记，在入住页面中填写老人基本信息、家属信息、入住合同信息、签约医生信息等，即可完成入住登记。')

add_heading_styled('2.3 床位管理模块', level=2)
doc.add_paragraph('床位管理包含三个功能：房间管理、床位预览、智能床位：')
doc.add_paragraph('房间管理：在房间管理页面维护养老院中的房间类型，包括房间图片、房间类型、床位配置、房间介绍等。')
doc.add_paragraph('床位预览：在床位预览页面维护养老院中的楼层、楼层下的房间、房间中的床位。')
doc.add_paragraph('智能床位：智能床位大屏展示养老院内所有智能床位分布、楼层和房间信息以及老人和设备信息、统计柱状图展示。')

add_heading_styled('2.4 护理管理模块', level=2)
doc.add_paragraph('护理管理包含四个功能：护理项目、护理计划、护理等级、护理人员：')
doc.add_paragraph('护理项目：在护理项目页面维护养老院能提供的护理项目（如洗头、洗脚、送餐等）。')
doc.add_paragraph('护理计划：在护理计划页面维护养老院中的护理计划，一个护理计划中可以关联多个护理项目。')
doc.add_paragraph('护理等级：在护理等级页面维护养老院中的护理等级，一个护理等级唯一对应一个护理计划。')
doc.add_paragraph('护理人员：在护理人员页面展示养老院中所有楼层，以及楼层下的房间和床位信息，可以维护某一床位上的护理员。')

add_heading_styled('2.5 智能监测模块', level=2)
doc.add_paragraph('智能监测包含三个功能：设备管理、告警规则、告警数据：')
doc.add_paragraph('设备管理：在设备管理页面维护对接第三方IOT平台的设备，包括设备在系统中的位置关联。')
doc.add_paragraph('告警规则：在告警规则页面维护系统中的告警规则，告警处理方案配置。')
doc.add_paragraph('告警数据：在告警数据页面查看所有上报的告警数据，并能对告警数据进行处理。')

add_heading_styled('2.6 微信小程序模块', level=2)
doc.add_paragraph('微信小程序包含三个功能：小程序登录、绑定家属、参观预约：')
doc.add_paragraph('小程序登录：微信小程序获取用户授权后，调用微信开放平台提供的登录接口进行登录。')
doc.add_paragraph('绑定家属：在微信小程序端绑定养老院已入住的家属信息，方便发起探访预约和查看老人的身体健康数据等。')
doc.add_paragraph('参观预约：家属在微信小程序端发起参观预约申请，按预约的时间到养老院进行参观。')

# ========== 三、项目职责 ==========
add_heading_styled('三、测试工作职责', level=1)

doc.add_paragraph('在项目中，我主要承担软件测试工程师的职责，负责全流程测试工作：')

responsibilities = [
    '测试计划制定：制定模块级测试计划，明确测试范围、测试策略、资源安排和交付标准',
    '测试用例设计：结合等价类划分法、边界值分析法、场景法设计覆盖全面的测试用例',
    '功能测试执行：执行Web端功能测试，覆盖入住管理、护理管理、床位管理、智能监测等核心模块',
    '接口测试：使用Apifox对后端API进行接口测试，包括参数校验、业务逻辑验证、异常场景覆盖',
    'Web自动化测试：使用Python + Selenium搭建自动化测试框架，覆盖核心业务流程',
    '缺陷管理：使用禅道对缺陷进行全生命周期管理，包括提交、跟踪、回归验证',
    '测试报告：输出测试报告，汇总测试结果和缺陷分析，评估版本质量',
    'AI功能测试：测试百度千帆大模型对接功能，验证体检报告分析结果的准确性和完整性',
    'IoT设备测试：配合测试智能监测模块的设备数据上报、告警规则触发、告警通知等',
    '跨团队协作：与产品经理、开发工程师、UI设计师高效协作，推动问题解决',
]

for r in responsibilities:
    doc.add_paragraph(r, style='List Bullet')

# ========== 四、项目重点难点 ==========
add_heading_styled('四、项目重点难点', level=1)

add_heading_styled('4.1 AI体检报告分析', level=2)
doc.add_paragraph(
    '该功能使用百度千帆AI模型服务解析老人体检报告。最初方案是上传PDF到OSS，让AI读取URL，'
    '但发现AI读取PDF能力有限。优化方案：使用工具读取PDF文件内容，拼接到Prompt中传给AI，'
    '虽然Token成本增加（约1元/次），但准确率大幅提升。'
)

add_heading_styled('4.2 设备数据高并发处理', level=2)
doc.add_paragraph(
    '智能监测模块中，IoT设备数据采集与上报是核心难点。设备通过MQTT上报数据到IOT平台，'
    'IOT平台通过RabbitMQ转发数据到后端。由于数据量大（单日约70万条），后端使用线程池异步处理，'
    '先写入Redis缓存，再批量落库。并用定时任务定期清理历史数据。'
)

add_heading_styled('4.3 告警机制设计', level=2)
doc.add_paragraph(
    '告警模块检测异常数据（老人体征异常、设备异常），根据不同设备类型通知不同负责人。'
    '使用Redis进行告警去重和静默：每次记录异常时检查Redis中是否已存在静默标记，'
    '避免重复告警。静默值在一段时间后自动失效，设备可再次告警。'
)

# ========== 五、测试效果 ==========
add_heading_styled('五、测试效果与产出', level=1)

stats = [
    ['编写测试用例', '120+ 条，覆盖所有核心业务模块'],
    ['执行功能测试', '8 轮迭代测试'],
    ['执行接口测试', '50+ 个接口，200+ 测试场景'],
    ['提交缺陷', '45 个，其中严重缺陷 8 个'],
    ['自动化脚本', '20+ 个自动化测试用例'],
    ['缺陷修复率', '95% 以上'],
    ['测试报告', '6 份版本测试报告'],
]
for k, v in stats:
    p = doc.add_paragraph()
    add_bold_text(p, f'{k}：')
    p.add_run(v)

doc.add_paragraph()
doc.add_paragraph(
    '通过系统的测试工作，确保了中州养老管理系统v1.0版本的顺利上线，上线后系统运行稳定，'
    '核心功能无P0/P1级别线上缺陷，获得了客户的好评。'
)

# ========== 六、个人总结 ==========
add_heading_styled('六、个人总结与收获', level=1)
doc.add_paragraph(
    '通过参与中州养老管理系统项目，我全面实践了软件测试的全流程，从需求评审、测试计划、'
    '用例设计到测试执行、缺陷管理和测试报告输出，积累了完整的项目测试经验。'
)
doc.add_paragraph('主要收获：')
gains = [
    '技术提升：熟练使用Python + Selenium进行Web自动化测试，使用Apifox进行接口测试',
    '工具掌握：熟练使用禅道进行缺陷管理，使用Git进行代码版本管理',
    '业务理解：深入理解养老行业的业务流程和需求特点',
    '团队协作：提升与产品、开发、UI等角色的沟通协作能力',
    '质量意识：建立全面的质量保障思维，从测试角度推动产品质量提升',
]
for g in gains:
    doc.add_paragraph(g, style='List Bullet')

# 保存
output_path = "C:/Users/ASUS/Desktop/生产实习/zyyl/中州养老管理系统-项目总结.docx"
doc.save(output_path)
print(f"项目总结已生成: {output_path}")
