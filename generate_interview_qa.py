"""
生成《慧享养老服务平台 - 测试面试题集》DOCX
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def add_q(text, level=2):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    return p

# 封面
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('慧享养老服务平台')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('软件测试面试题集')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 102, 153)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('基于慧享养老服务平台项目的面试题及答案\n涵盖：项目经验 / 测试理论 / 工具实战 / 技术深度')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(102, 102, 102)

doc.add_page_break()

# ===== 面试题 =====

add_q('一、项目经验类', level=1)

add_q('1. 请介绍一下你最近做的项目')
doc.add_paragraph(
    '我最近参与的是慧享养老服务平台，这是一个专门为养老机构定制开发的养老服务管理平台。'
    '我主要负责后台管理系统的测试工作。\n\n'
    '后台管理系统供养老院员工使用，主要包含入住管理、退住管理、护理管理（护理等级、护理计划、护理项目）、'
    '床位管理（楼层、房间、床位）、智能监测（设备管理、告警规则、告警数据）等核心模块。\n\n'
    '技术方面：后端使用Spring Boot 3.5.x + MyBatis + Redis，前端使用Vue3 + Element Plus，'
    '集成了百度千帆大模型做AI体检报告分析。\n\n'
    '我在项目中担任软件测试工程师，负责全流程测试工作。'
)

add_q('2. 你在项目中主要负责什么？')
doc.add_paragraph(
    '我主要负责以下测试工作：\n\n'
    '1）测试计划制定：制定模块级测试计划，明确范围、策略和交付标准\n'
    '2）测试用例设计：结合等价类划分、边界值分析、场景法设计全面用例\n'
    '3）功能测试执行：覆盖入住管理、护理管理、床位管理、智能监测等核心模块\n'
    '4）接口测试：使用Apifox对后端50+接口进行参数校验和业务场景验证\n'
    '5）Web自动化测试：使用Python + Selenium + pytest搭建自动化测试框架\n'
    '6）缺陷管理：使用禅道进行缺陷全生命周期管理\n'
    '7）AI功能测试：测试百度千帆大模型对接功能\n'
    '8）测试报告：每次迭代输出质量报告'
)

add_q('3. 项目中遇到的比较有挑战的问题是什么？')
doc.add_paragraph(
    '挑战最大的有两个：\n\n'
    '第一是AI体检报告分析功能的测试。老人体检报告是PDF格式，需要上传后调用AI模型分析。'
    '最初方案是传PDF链接给AI，但发现AI读取PDF内容有限。后改为先解析PDF文本再传给AI，'
    'Token成本增加了但准确率大幅提升。测试时我需要准备多种体检报告样本，'
    '验证AI分析结果的准确性、格式规范性，以及异常情况的处理。\n\n'
    '第二是IoT设备监测模块的测试。设备通过MQTT上报数据到IOT平台，再通过RabbitMQ转发到后端。'
    '单日数据量约70万条，后端使用线程池异步处理。测试时需要模拟设备数据上报、'
    '验证告警规则触发、告警去重、静默机制等。'
)

add_q('4. 养老系统有哪些核心业务？')
doc.add_paragraph(
    '养老院的核心业务流程如下：\n\n'
    '护理等级（等级划分）→ 护理项目（具体服务：洗头、洗脚、送餐等）\n'
    '→ 护理计划（按等级配置服务组合）→ 床位分配（楼层→房间→床位）\n'
    '→ 入住登记（录入老人信息、家属信息、合同信息）\n'
    '→ 享受护理服务（按等级执行护理计划）\n'
    '→ 智能监测（健康数据采集、告警通知）\n\n'
    '退住流程：家属申请 → 费用结算 → 物品清点 → 退住确认'
)

add_q('5. 说一下入住和退住的完整流程')
doc.add_paragraph(
    '入住流程：\n'
    '1）咨询参观：家属通过电话/网站了解养老院，预约实地参观\n'
    '2）健康评估：专业医护人员对老人进行健康评估，确定护理等级\n'
    '3）签约缴费：签订入住协议，缴纳押金等费用\n'
    '4）分配床位：根据评估结果分配房间和床位，安装智能设备\n'
    '5）正式入住：老人入住，工作人员介绍制度和活动安排\n\n'
    '退住流程：\n'
    '1）提交申请：家属提出退住申请\n'
    '2）费用核算：核算入住期间所有费用，退还押金\n'
    '3）物品清点：清点老人个人物品，检查房间设施\n'
    '4）办理退住：填写退住确认单，完成退住手续\n'
    '5）离院：老人离开养老院，工作人员欢送'
)

# ===== 第二部分：测试理论 =====

add_q('二、测试理论类', level=1)

add_q('6. 你之前公司的测试流程是什么样的？')
doc.add_paragraph(
    '测试流程包括以下阶段：\n'
    '1）需求评审：产品经理发起需求评审，测试提前了解需求并提出疑问\n'
    '2）测试计划：制定测试计划，明确测试范围、策略、资源、进度\n'
    '3）用例设计：使用多种方法设计测试用例，并进行用例评审\n'
    '4）测试执行：开发提测后执行功能测试、接口测试、兼容性测试\n'
    '5）缺陷管理：发现Bug提交到禅道，跟踪开发修复，回归验证\n'
    '6）测试报告：输出测试报告，评估版本质量\n'
    '7）上线跟踪：协助上线，跟进线上问题'
)

add_q('7. 如何设计测试用例？举例说明')
doc.add_paragraph(
    '我综合使用多种方法设计用例：\n\n'
    '以入住登记模块的"年龄"字段为例：\n\n'
    '1）等价类划分：有效等价类（18-150岁），无效等价类（<18岁，>150岁，非数字，空值）\n'
    '2）边界值分析：测试17、18、150、151、-1、0、abc这些边界\n'
    '3）场景法：基本流（填写完整信息 → 保存成功）、备选流（未填必填项 → 提示错误）\n'
    '4）错误推测法：特殊字符（<script>、SQL注入）、超长字段（姓名50字）、重复提交\n'
    '最终输出Excel格式的测试用例，包含编号、名称、前置条件、步骤、预期结果、优先级'
)

add_q('8. 如何区分前端Bug还是后端Bug？')
doc.add_paragraph(
    '定位步骤如下：\n'
    '1）打开浏览器F12开发者工具\n'
    '2）查看Network请求：\n'
    '   - 请求未发出 → 前端问题\n'
    '   - 请求参数错误 → 前端问题\n'
    '   - 请求正常但响应数据错误 → 后端问题\n'
    '   - 响应数据正确但页面显示错误 → 前端问题\n'
    '3）查看控制台Console是否有JS报错（前端问题）\n'
    '4）查看后端日志（tail -f / ELK）检查是否有异常（后端问题）\n'
    '5）接口测试：直接用Apifox调用接口，如果接口返回正确则是前端问题'
)

add_q('9. 接口测试主要测什么？')
doc.add_paragraph(
    '接口测试主要从以下几个维度进行：\n'
    '1）功能验证：接口是否按照文档返回正确数据\n'
    '2）参数校验：必填项、参数类型、参数长度、参数格式\n'
    '3）业务逻辑：正业务流程、异常流程、边界条件\n'
    '4）安全性：权限校验、SQL注入、XSS、未授权访问\n'
    '5）性能：响应时间是否在合理范围内\n'
    '6）幂等性：多次调用结果是否一致\n'
    '7）异常场景：超时、网络中断、数据不存在'
)

add_q('10. 如何进行回归测试？')
doc.add_paragraph(
    '回归测试策略：\n'
    '1）确定回归范围：分析代码变更影响的范围\n'
    '2）选择回归用例：\n'
    '   - 受影响的模块全部用例\n'
    '   - 核心业务流程（P0用例必跑）\n'
    '   - 之前出现过Bug的用例\n'
    '   - 与修改代码相关的上下游模块用例\n'
    '3）优先使用自动化执行回归\n'
    '4）回归直到无新增Bug且所有P0/P1用例通过\n'
    '5）输出回归测试报告，评估风险'
)

add_q('11. 如果时间不够，你会怎么安排测试？')
doc.add_paragraph(
    '当测试时间不足时，我会按优先级安排：\n'
    '1）首先保证P0用例（核心业务流程）全部覆盖\n'
    '2）保证P1用例（主要功能）覆盖\n'
    '3）P2/P3用例选择性覆盖（可适当放弃边缘场景）\n'
    '4）优先使用自动化测试提高效率\n'
    '5）风险沟通：及时向项目经理反馈测试风险，明确未覆盖的范围\n'
    '6）加班保障上线质量'
)

# ===== 第三部分：工具实战 =====

add_q('三、测试工具实战类', level=1)

add_q('12. Selenium如何处理动态加载的元素？')
doc.add_paragraph(
    '使用显式等待（WebDriverWait + expected_conditions）来处理动态元素：\n\n'
    'from selenium.webdriver.support.ui import WebDriverWait\n'
    'from selenium.webdriver.support import expected_conditions as EC\n\n'
    '# 等待元素出现\n'
    'element = WebDriverWait(driver, 10).until(\n'
    '    EC.presence_of_element_located((By.ID, "dynamicElement"))\n'
    ')\n\n'
    '# 对于Element Plus的Dialog等异步组件，等待其完全渲染\n'
    'WebDriverWait(driver, 10).until(\n'
    '    EC.visibility_of_element_located((By.XPATH, "//div[@class=\'el-dialog\']"))\n'
    ')'
)

add_q('13. Apifox中如何设置token自动管理？')
doc.add_paragraph(
    '在Apifox的环境变量中设置token自动管理：\n'
    '1）在"环境管理"中配置base_url\n'
    '2）编写前置脚本，自动登录获取token：\n\n'
    'const token = pm.environment.get("token");\n'
    'if (!token) {\n'
    '    pm.sendRequest({\n'
    '        url: pm.environment.get("base_url") + "/login",\n'
    '        method: "POST",\n'
    '        header: {"Content-Type": "application/json"},\n'
    '        body: { mode: "raw", raw: JSON.stringify({username:"admin", password:"admin123", code:"1", uuid:"test"}) }\n'
    '    }, function(err, res) {\n'
    '        pm.environment.set("token", res.json().token);\n'
    '    });\n'
    '}'
)

add_q('14. 禅道中缺陷的生命周期是怎样的？')
doc.add_paragraph(
    '完整生命周期：\n'
    '新建(测试提Bug) → 已确认(开发确认) → 已解决(开发修复) → 已关闭(验证通过)\n\n'
    '分支流程：\n'
    '- 已解决 → 激活（回归不通过，重新激活）\n'
    '- 已确认 → 不予解决（设计如此，非Bug）\n'
    '- 已确认 → 延期处理（下版本修复）\n'
    '- 已解决 → 已关闭 → 激活（线上重现，重新激活）'
)

add_q('15. 如何用Python + Selenium实现数据驱动测试？')
doc.add_paragraph(
    '使用pytest的parametrize装饰器实现数据驱动：\n\n'
    '@pytest.mark.parametrize("username,password,expected", [\n'
    '    ("admin", "admin123", True),   # 正确账号密码\n'
    '    ("admin", "wrong", False),     # 错误密码\n'
    '    ("", "admin123", False),       # 空用户名\n'
    '    ("admin", "", False),          # 空密码\n'
    '    ("admin@!", "admin123", False),# 特殊字符\n'
    '])\n'
    'def test_login(self, login_page, username, password, expected):\n'
    '    login_page.login(username, password)\n'
    '    assert login_page.is_login_success() == expected\n\n'
    '也可以从Excel或JSON文件读取测试数据，用pytest的fixture动态注入。'
)

add_q('16. Apifox中如何做场景测试（关联接口）？')
doc.add_paragraph(
    '在Apifox中创建"测试场景"，将多个接口按业务流程串联：\n'
    '1）创建场景"入住管理完整流程"\n'
    '2）添加步骤1：POST /login（登录获取token）\n'
    '3）添加步骤2：GET /nursing/floor/list（获取楼层）\n'
    '4）添加步骤3：POST /nursing/checkIn（新增入住，依赖前两步数据）\n'
    '5）设置变量传递：将login返回的token设为全局变量\n'
    '6）设置断言：每个步骤都设置相应的断言\n'
    '7）运行场景：一键执行整个业务流程\n'
    '8）查看报告：检查每个步骤的执行结果'
)

# ===== 第四部分：技术深度 =====

add_q('四、技术深度类', level=1)

add_q('17. 说说你对自动化测试框架的理解')
doc.add_paragraph(
    '一个成熟的自动化测试框架应该包含：\n\n'
    '1）用例管理：支持数据驱动、关键字驱动\n'
    '2）定位策略：统一元素定位方式，支持多种定位器\n'
    '3）等待机制：合理使用显式/隐式等待\n'
    '4）报告生成：生成HTML格式的测试报告\n'
    '5）日志记录：记录测试执行日志，方便定位问题\n'
    '6）截图功能：失败时自动截图\n'
    '7）数据分离：测试数据和代码分离，便于维护\n'
    '8）持续集成：集成到Jenkins自动执行\n'
    '9）Page Object模式：页面和操作分层\n'
    '10）异常处理：完善的异常捕获和重试机制'
)

add_q('18. 什么是PO模式？用完有什么好处？')
doc.add_paragraph(
    'PO（Page Object）模式是Web自动化的设计模式：\n\n'
    '核心思想：将每个页面封装成一个类\n'
    '- 类中定义该页面的所有元素定位\n'
    '- 类中封装该页面的所有操作\n\n'
    '优点：\n'
    '- 维护性高：页面元素变化只需修改一处\n'
    '- 代码复用：多个测试用例可复用页面操作\n'
    '- 可读性好：测试用例关注业务，不关注元素细节\n'
    '- 减少耦合：页面变化不影响测试用例\n\n'
    '示例：LoginPage类封装了登录页面的所有元素和操作，'
    'test_login.py直接调用login_page.login()即可。'
)

add_q('19. SQL慢查询如何优化？')
doc.add_paragraph(
    '优化步骤：\n'
    '1）开启慢查询日志，定位慢SQL\n'
    '2）使用EXPLAIN查看执行计划（关注type、key、key_len、rows、Extra）\n'
    '3）优化手段：\n'
    '   - 合理添加索引（注意联合索引最左匹配原则）\n'
    '   - 避免SELECT *，只查需要的字段\n'
    '   - 避免WHERE中使用函数、隐式类型转换\n'
    '   - 避免前导模糊查询（LIKE "%xxx" 索引失效）\n'
    '   - 大表考虑分库分表\n'
    '   - 使用覆盖索引（Extra显示Using index）\n'
    '4）验证优化效果，对比优化前后的执行时间'
)

add_q('20. Redis在项目中怎么用的？')
doc.add_paragraph(
    '在慧享养老系统中，Redis的应用场景：\n'
    '1）缓存：缓存高频查询数据（用户信息、护理等级、床位状态），减少数据库压力\n'
    '2）缓存一致性策略：写操作时删除缓存，下次查询重新加载\n'
    '3）验证码：临时存储验证码，设置过期时间\n'
    '4）登录Token：JWT Token的校验\n'
    '5）告警静默：使用Redis存储告警静默标记，避免重复告警\n'
    '6）设备数据：IoT设备上报数据先写入Redis，再批量落库\n'
    '7）定时任务：分布式锁防止定时任务重复执行'
)

# 保存
output_path = "C:/Users/ASUS/Desktop/生产实习/zyyl/慧享养老服务平台-面试题集.docx"
doc.save(output_path)
print(f"面试题集已生成: {output_path}")
