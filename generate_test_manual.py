"""
生成《慧享养老服务平台-测试全流程操作手册》
覆盖：Python基础、Selenium Web自动化、Apifox接口测试、禅道缺陷管理
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ========== 样式设置 ==========
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
    return h


def add_code_block(text):
    """添加代码块样式"""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.5 + level * 0.5)
    return p


def add_table(headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    return table


# =============================================
# 封面
# =============================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('慧享养老服务平台')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 51, 102)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('测试全流程操作手册')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0, 102, 153)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(
    '涵盖工具：Python基础 · Selenium Web自动化 · Apifox接口测试 · 禅道缺陷管理\n'
    '版本：v1.0\n'
    '日期：2024年'
)
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(102, 102, 102)

doc.add_page_break()

# =============================================
# 目录页
# =============================================
add_heading('目 录', level=1)
toc_items = [
    '第一章  Python基础与测试应用',
    '  1.1 Python环境搭建',
    '  1.2 基础语法速览',
    '  1.3 常用测试库',
    '  1.4 项目实战：养老系统测试数据生成',
    '第二章  Selenium Web自动化测试',
    '  2.1 Selenium简介与环境搭建',
    '  2.2 元素定位策略',
    '  2.3 常用操作',
    '  2.4 等待策略',
    '  2.5 测试框架搭建',
    '  2.6 项目实战：养老系统Web自动化',
    '第三章  Apifox接口测试',
    '  3.1 Apifox简介',
    '  3.2 接口文档管理',
    '  3.3 接口调试与用例管理',
    '  3.4 自动化测试',
    '  3.5 项目实战：养老系统接口测试',
    '第四章  禅道缺陷管理',
    '  4.1 禅道简介',
    '  4.2 缺陷生命周期',
    '  4.3 缺陷提交规范',
    '  4.4 项目实战：养老系统缺陷管理',
    '第五章  测试全流程实战',
    '  5.1 测试计划',
    '  5.2 测试用例设计',
    '  5.3 测试执行',
    '  5.4 缺陷跟踪',
    '  5.5 测试报告',
    '附录：常见面试题汇总'
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# =============================================
# 第一章 Python基础与测试应用
# =============================================
add_heading('第一章  Python基础与测试应用', level=1)

add_heading('1.1 Python环境搭建', level=2)
doc.add_paragraph(
    'Python是测试工程师必备的编程语言，在接口测试、Web自动化、数据处理等方面有广泛应用。'
)

add_heading('1.1.1 安装Python', level=3)
doc.add_paragraph('步骤一：访问 Python官网 (https://www.python.org/downloads/) 下载Python 3.9+版本')
doc.add_paragraph('步骤二：安装时勾选"Add Python to PATH"')
doc.add_paragraph('步骤三：验证安装')
add_code_block('python --version\npip --version')

add_heading('1.1.2 虚拟环境配置', level=3)
doc.add_paragraph('为每个项目创建独立的虚拟环境，避免依赖冲突：')
add_code_block(
    '# 创建虚拟环境\n'
    'python -m venv zzyl_test_env\n\n'
    '# 激活虚拟环境（Windows）\n'
    'zzyl_test_env\\Scripts\\activate\n\n'
    '# 激活虚拟环境（Mac/Linux）\n'
    'source zzyl_test_env/bin/activate'
)

add_heading('1.1.3 安装测试相关依赖', level=3)
add_code_block(
    '# 创建 requirements.txt\n'
    'selenium==4.30.0\n'
    'pytest==8.3.5\n'
    'pytest-html==4.1.1\n'
    'requests==2.32.3\n'
    'openpyxl==3.1.5\n'
    'python-dotenv==1.1.0\n\n'
    '# 一键安装\n'
    'pip install -r requirements.txt'
)

add_heading('1.2 基础语法速览', level=2)

add_heading('1.2.1 变量与数据类型', level=3)
add_code_block(
    '# 常用数据类型\n'
    'name = "慧享养老"          # 字符串 str\n'
    'version = 1.0               # 浮点数 float\n'
    'total_cases = 100           # 整数 int\n'
    'is_passed = True            # 布尔值 bool\n'
    'test_data = ["TC001", "TC002", "TC003"]  # 列表 list\n'
    'test_result = {"TC001": "通过", "TC002": "失败"}  # 字典 dict\n\n'
    '# print 调试\n'
    'print(f"测试用例总数: {total_cases}")  # f-string格式化'
)

add_heading('1.2.2 条件与循环', level=3)
add_code_block(
    '# if-elif-else 判断\n'
    'status = "通过"\n'
    'if status == "通过":\n'
    '    print("测试通过")\n'
    'elif status == "失败":\n'
    '    print("测试失败")\n'
    'else:\n'
    '    print("其他状态")\n\n'
    '# for 循环遍历\n'
    'test_cases = ["TC001", "TC002", "TC003"]\n'
    'for case in test_cases:\n'
    '    print(f"执行用例: {case}")\n\n'
    '# while 循环\n'
    'retry_count = 0\n'
    'while retry_count < 3:\n'
    '    print(f"第 {retry_count+1} 次重试")\n'
    '    retry_count += 1'
)

add_heading('1.2.3 函数与异常处理', level=3)
add_code_block(
    'def run_test(case_id, expected=True):\n'
    '    """执行单个测试用例"""\n'
    '    try:\n'
    '        result = execute_case(case_id)\n'
    '        assert result == expected, f"用例{case_id}执行失败"\n'
    '        return True\n'
    '    except AssertionError as e:\n'
    '        print(f"断言失败: {e}")\n'
    '        return False\n'
    '    except Exception as e:\n'
    '        print(f"执行异常: {e}")\n'
    '        return False\n'
    '    finally:\n'
    '        print(f"用例{case_id}执行完毕")'
)

add_heading('1.3 常用测试库', level=2)

test_libs = [
    ['requests', 'HTTP库', '接口测试、API调用', 'import requests'],
    ['selenium', '浏览器自动化', 'Web功能自动化测试', 'from selenium import webdriver'],
    ['pytest', '测试框架', '用例组织、断言、报告', 'import pytest'],
    ['openpyxl', 'Excel操作', '测试数据管理', 'from openpyxl import load_workbook'],
    ['pandas', '数据分析', '测试结果分析', 'import pandas as pd'],
    ['json', 'JSON处理', 'API请求/响应处理', 'import json'],
]
add_table(['库名', '类型', '用途', '导入方式'], test_libs)

add_heading('1.4 项目实战：养老系统测试数据生成', level=2)
doc.add_paragraph('以慧享养老服务平台为例，生成测试数据：')
add_code_block(
    'import json\n'
    'import random\n\n\n'
    'def generate_elder_test_data(count=10):\n'
    '    """生成老人入住测试数据"""\n'
    '    names = ["张建国", "李明华", "王德福", "刘秀英", "陈志强"]\n'
    '    levels = ["半自理", "全护理", "失智照护", "自理"]\n'
    '    rooms = [f"{floor}0{room}" for floor in range(1, 6) for room in range(1, 5)]\n\n'
    '    test_data = []\n'
    '    for i in range(count):\n'
    '        elder = {\n'
    '            "name": random.choice(names),\n'
    '            "age": random.randint(65, 95),\n'
    '            "gender": random.choice(["男", "女"]),\n'
    '            "phone": f"1{random.randint(30, 99)}" + "".join(\n'
    '                [str(random.randint(0, 9)) for _ in range(8)]\n'
    '            ),\n'
    '            "nursingLevel": random.choice(levels),\n'
    '            "roomNo": random.choice(rooms),\n'
    '            "checkinDate": "2024-06-01",\n'
    '            "emergencyContact": "子女",\n'
    '            "emergencyPhone": "13800138000"\n'
    '        }\n'
    '        test_data.append(elder)\n\n'
    '    return json.dumps(test_data, ensure_ascii=False, indent=2)\n\n\n'
    'if __name__ == "__main__":\n'
    '    data = generate_elder_test_data(5)\n'
    '    print(data)'
)

doc.add_page_break()

# =============================================
# 第二章 Selenium Web自动化测试
# =============================================
add_heading('第二章  Selenium Web自动化测试', level=1)

add_heading('2.1 Selenium简介与环境搭建', level=2)
doc.add_paragraph(
    'Selenium是一个用于Web应用程序测试的自动化工具，支持多种浏览器（Chrome、Firefox、Edge等）。'
    '它直接运行在浏览器中，模拟用户操作，实现自动化测试。'
)

add_heading('2.1.1 安装Selenium', level=3)
add_code_block('pip install selenium')

add_heading('2.1.2 ChromeDriver配置', level=3)
doc.add_paragraph('方式一：手动下载ChromeDriver')
doc.add_paragraph('1. 查看Chrome版本：chrome://settings/help')
doc.add_paragraph('2. 访问 https://chromedriver.chromium.org/ 下载对应版本')
doc.add_paragraph('3. 将chromedriver.exe放入PATH目录或项目根目录')

doc.add_paragraph('方式二：使用WebDriver Manager（推荐）')
add_code_block('# 自动管理浏览器驱动\nfrom webdriver_manager.chrome import ChromeDriverManager\nfrom selenium.webdriver.chrome.service import Service\n\nservice = Service(ChromeDriverManager().install())\ndriver = webdriver.Chrome(service=service)')

add_heading('2.1.3 浏览器基本操作', level=3)
add_code_block(
    'from selenium import webdriver\nfrom selenium.webdriver.chrome.options import Options\n\n'
    '# 创建Chrome浏览器实例\nchrome_options = Options()\n'
    'chrome_options.add_argument("--window-size=1920,1080")\n'
    'chrome_options.add_argument("--disable-notifications")\n\n'
    'driver = webdriver.Chrome(options=chrome_options)\n\n'
    '# 打开网页\ndriver.get("http://localhost:80")\n\n'
    '# 获取页面标题\nprint(driver.title)\n\n'
    '# 截图\ndriver.save_screenshot("homepage.png")\n\n'
    '# 关闭浏览器\ndriver.quit()'
)

add_heading('2.2 元素定位策略', level=2)
doc.add_paragraph('Selenium提供了8种元素定位方式：')

locators = [
    ['ID', 'id', "driver.find_element(By.ID, 'username')"],
    ['Name', 'name', "driver.find_element(By.NAME, 'password')"],
    ['Class Name', 'class name', "driver.find_element(By.CLASS_NAME, 'el-input')"],
    ['Tag Name', 'tag name', "driver.find_element(By.TAG_NAME, 'button')"],
    ['Link Text', 'link text', "driver.find_element(By.LINK_TEXT, '登录')"],
    ['Partial Link Text', 'partial link', "driver.find_element(By.PARTIAL_LINK_TEXT, '登录')"],
    ['XPath', 'xpath', "driver.find_element(By.XPATH, '//input[@placeholder=\"用户名\"]')"],
    ['CSS Selector', 'css selector', "driver.find_element(By.CSS_SELECTOR, '#app > div')"],
]
add_table(['定位方式', 'By类型', '示例'], locators)

doc.add_paragraph()
doc.add_paragraph('推荐优先使用ID定位，其次是XPath和CSS Selector。'
                   '对于Element Plus/Vue项目，常用XPath定位带placeholder的元素。')

add_heading('2.2.1 XPath定位详解', level=3)
add_code_block(
    '# 常用XPath语法\n\n'
    '# 根据文本定位\n'
    'driver.find_element(By.XPATH, "//span[contains(text(), \'入住管理\')]")\n\n'
    '# 根据属性定位\n'
    'driver.find_element(By.XPATH, "//input[@placeholder=\'用户名\']")\n\n'
    '# 组合定位\n'
    'driver.find_element(By.XPATH, "//button[contains(@class, \'login-btn\')]")\n\n'
    '# 层级定位（查找表单下的提交按钮）\n'
    'driver.find_element(By.XPATH, "//form//button[@type=\'submit\']")\n\n'
    '# 使用索引定位\n'
    'driver.find_element(By.XPATH, "(//div[@class=\'menu\']//span)[2]")'
)

add_heading('2.3 常用操作', level=2)
add_code_block(
    '# 输入文本\n'
    'driver.find_element(By.ID, "username").send_keys("admin")\n\n'
    '# 点击元素\n'
    'driver.find_element(By.XPATH, "//button[contains(text(), \'登录\')]").click()\n\n'
    '# 获取文本\n'
    'text = driver.find_element(By.CLASS_NAME, "title").text\n\n'
    '# 获取属性\n'
    'value = driver.find_element(By.ID, "username").get_attribute("value")\n\n'
    '# 清除文本\n'
    'driver.find_element(By.ID, "username").clear()\n\n'
    '# 判断元素是否可见\n'
    'is_displayed = driver.find_element(By.ID, "submit").is_displayed()\n\n'
    '# 下拉选择\n'
    'from selenium.webdriver.support.ui import Select\n'
    'select = Select(driver.find_element(By.NAME, "status"))\n'
    'select.select_by_visible_text("通过")\n'
    '# select.select_by_value("1")\n'
    '# select.select_by_index(0)\n\n'
    '# 执行JavaScript\n'
    'driver.execute_script("window.scrollTo(0, document.body.scrollHeight)")\n'
    'driver.execute_script("arguments[0].click();", element)\n\n'
    '# 切换frame/iframe\n'
    'driver.switch_to.frame("iframe_id")\n'
    'driver.switch_to.default_content()  # 切回主文档\n\n'
    '# 弹窗处理\n'
    'alert = driver.switch_to.alert\n'
    'alert.accept()  # 确定\n'
    'alert.dismiss()  # 取消'
)

add_heading('2.4 等待策略', level=2)
doc.add_paragraph('等待是Web自动化中最关键的技术之一，用于处理页面加载和异步渲染。')

add_heading('2.4.1 隐式等待', level=3)
add_code_block(
    'from selenium import webdriver\n\n'
    'driver = webdriver.Chrome()\n'
    '# 设置隐式等待(全局)，在查找元素时等待最多10秒\n'
    'driver.implicitly_wait(10)'
)

add_heading('2.4.2 显式等待', level=3)
add_code_block(
    'from selenium.webdriver.support.ui import WebDriverWait\n'
    'from selenium.webdriver.support import expected_conditions as EC\n\n'
    '# 等待元素出现（推荐）\n'
    'element = WebDriverWait(driver, 10).until(\n'
    '    EC.presence_of_element_located((By.ID, "myElement"))\n'
    ')\n\n'
    '# 等待元素可点击\n'
    'element = WebDriverWait(driver, 10).until(\n'
    '    EC.element_to_be_clickable((By.XPATH, "//button[text()=\'登录\']"))\n'
    ')\n\n'
    '# 等待元素可见\n'
    'WebDriverWait(driver, 10).until(\n'
    '    EC.visibility_of_element_located((By.ID, "result"))\n'
    ')\n\n'
    '# 等待元素不可见\n'
    'WebDriverWait(driver, 10).until(\n'
    '    EC.invisibility_of_element_located((By.ID, "loading"))\n'
    ')\n\n'
    '# 等待特定文本出现\n'
    'WebDriverWait(driver, 10).until(\n'
    '    EC.text_to_be_present_in_element((By.ID, "status"), "通过")\n'
    ')\n\n'
    '# 常用expected_conditions\n'
    '# EC.alert_is_present() - 等待弹窗\n'
    '# EC.title_contains("xxx") - 等待标题包含xxx\n'
    '# EC.frame_to_be_available_and_switch_to_it() - 等待frame\n'
    '# EC.element_to_be_selected() - 等待元素被选中'
)

add_heading('2.4.3 强制等待', level=3)
add_code_block(
    'import time\n'
    'time.sleep(2)  # 固定等待2秒（不推荐，影响效率）'
)

add_heading('2.5 测试框架搭建', level=2)
doc.add_paragraph('推荐使用 pytest + Selenium 搭建自动化测试框架：')

add_heading('2.5.1 项目结构', level=3)
add_code_block(
    'zzyl-autotest/\n'
    '├── conftest.py          # pytest配置和fixture\n'
    '├── requirements.txt     # 依赖管理\n'
    '├── pages/               # Page Object模式\n'
    '│   ├── login_page.py\n'
    '│   ├── checkin_page.py\n'
    '│   └── base_page.py\n'
    '├── testcases/           # 测试用例\n'
    '│   ├── test_login.py\n'
    '│   ├── test_checkin.py\n'
    '│   └── test_nursing.py\n'
    '├── data/                # 测试数据\n'
    '│   └── test_data.xlsx\n'
    '├── reports/             # 测试报告\n'
    '│   └── report.html\n'
    '└── utils/               # 工具类\n'
    '    ├── excel_util.py\n'
    '    └── log_util.py'
)

add_heading('2.5.2 Page Object模式（PO模式）', level=3)
doc.add_paragraph('PO模式是Web自动化的最佳实践，将页面元素和操作封装在独立的类中：')
add_code_block(
    '# pages/login_page.py\n'
    'from selenium.webdriver.common.by import By\n\n\n'
    'class LoginPage:\n'
    '    """登录页面对象"""\n\n'
    '    # 元素定位\n'
    '    USERNAME_INPUT = (By.XPATH, "//input[@placeholder=\'用户名\']")\n'
    '    PASSWORD_INPUT = (By.XPATH, "//input[@placeholder=\'密码\']")\n'
    '    CAPTCHA_INPUT = (By.XPATH, "//input[@placeholder=\'验证码\']")\n'
    '    LOGIN_BUTTON = (By.XPATH, "//button[contains(@class, \'login-btn\')]")\n\n'
    '    def __init__(self, driver):\n'
    '        self.driver = driver\n\n'
    '    def input_username(self, username):\n'
    '        self.driver.find_element(*self.USERNAME_INPUT).send_keys(username)\n\n'
    '    def input_password(self, password):\n'
    '        self.driver.find_element(*self.PASSWORD_INPUT).send_keys(password)\n\n'
    '    def input_captcha(self, captcha):\n'
    '        self.driver.find_element(*self.CAPTCHA_INPUT).send_keys(captcha)\n\n'
    '    def click_login(self):\n'
    '        self.driver.find_element(*self.LOGIN_BUTTON).click()\n\n'
    '    def login(self, username, password, captcha="1"):\n'
    '        """封装完整的登录操作"""\n'
    '        self.input_username(username)\n'
    '        self.input_password(password)\n'
    '        self.input_captcha(captcha)\n'
    '        self.click_login()'
)

add_heading('2.6 项目实战：养老系统Web自动化', level=2)

add_heading('2.6.1 登录模块测试', level=3)
add_code_block(
    'import pytest\n'
    'from selenium import webdriver\n'
    'from selenium.webdriver.common.by import By\n'
    'from selenium.webdriver.support.ui import WebDriverWait\n'
    'from selenium.webdriver.support import expected_conditions as EC\n\n\n'
    'class TestLogin:\n'
    '    """慧享养老系统登录测试"""\n\n'
    '    def test_login_success(self, driver):\n'
    '        """TC-001: 正确账号密码登录"""\n'
    '        driver.get("http://localhost:80")\n\n'
    '        wait = WebDriverWait(driver, 20)\n'
    '        wait.until(EC.presence_of_element_located(\n'
    '            (By.XPATH, "//input[@placeholder=\'用户名\']")))\n\n'
    '        # 输入登录信息\n'
    '        driver.find_element(By.XPATH, "//input[@placeholder=\'用户名\']").send_keys("admin")\n'
    '        driver.find_element(By.XPATH, "//input[@placeholder=\'密码\']").send_keys("admin123")\n'
    '        driver.find_element(By.XPATH, "//input[@placeholder=\'验证码\']").send_keys("1")\n'
    '        driver.find_element(By.XPATH, "//button[contains(@class, \'login-btn\')]").click()\n\n'
    '        # 验证登录成功\n'
    '        try:\n'
    '            WebDriverWait(driver, 10).until(\n'
    '                EC.presence_of_element_located(\n'
    '                    (By.XPATH, "//div[contains(text(), \'慧享养老\')]")))\n'
    '            print("登录成功！")\n'
    '        except:\n'
    '            driver.save_screenshot("login_fail.png")\n'
    '            assert False, "登录失败"'
)

add_heading('2.6.2 入住管理测试', level=3)
add_code_block(
    'class TestCheckin:\n'
    '    """入住管理模块测试"""\n\n'
    '    def test_checkin_list(self, login_driver):\n'
    '        """入住列表加载测试"""\n'
    '        # 点击入住管理菜单\n'
    '        login_driver.find_element(\n'
    '            By.XPATH, "//span[contains(text(), \'入住管理\')]").click()\n\n'
    '        # 等待表格加载\n'
    '        WebDriverWait(login_driver, 10).until(\n'
    '            EC.presence_of_element_located((By.XPATH, "//table")))\n'
    '        print("入住列表加载成功")\n\n'
    '    def test_checkin_search_by_name(self, login_driver):\n'
    '        """按姓名搜索入住信息"""\n'
    '        login_driver.find_element(\n'
    '            By.XPATH, "//span[contains(text(), \'入住管理\')]").click()\n\n'
    '        # 输入搜索条件\n'
    '        search_input = login_driver.find_element(\n'
    '            By.XPATH, "//input[@placeholder=\'请输入老人姓名\']")\n'
    '        search_input.send_keys("张")\n'
    '        login_driver.find_element(\n'
    '            By.XPATH, "//button[contains(@class, \'el-button--primary\')]").click()\n\n'
    '        time.sleep(1)\n'
    '        # 验证搜索结果\n'
    '        rows = login_driver.find_elements(By.XPATH, "//table//tbody//tr")\n'
    '        print(f"搜索结果行数: {len(rows)}")'
)

doc.add_page_break()

# =============================================
# 第三章 Apifox接口测试
# =============================================
add_heading('第三章  Apifox接口测试', level=1)

add_heading('3.1 Apifox简介', level=2)
doc.add_paragraph(
    'Apifox是一款集API文档管理、API调试、API Mock、API自动化测试于一体的协作工具。'
    '它支持Swagger、OpenAPI等规范，可一键导入接口文档，极大提升接口测试效率。'
)

doc.add_paragraph('Apifox核心功能：')
features = [
    ['接口文档管理', '支持Swagger/OpenAPI导入导出，在线编辑'],
    ['接口调试', '支持GET/POST/PUT/DELETE等请求方式'],
    ['Mock服务', '根据接口定义自动生成Mock数据'],
    ['自动化测试', '支持场景测试、参数化、断言校验'],
    ['CI/CD集成', '支持命令行运行，集成Jenkins/GitLab'],
    ['团队协作', '支持多人在线协作，版本对比'],
]
add_table(['功能', '说明'], features)

add_heading('3.2 接口文档管理', level=2)

add_heading('3.2.1 导入接口文档', level=3)
doc.add_paragraph('慧享养老服务平台基于SpringDoc OpenAPI自动生成接口文档。')
doc.add_paragraph('导入方式：')
doc.add_paragraph('方式一：自动扫描（推荐）')
add_code_block(
    '1. 打开Apifox，创建项目"慧享养老服务平台"\n'
    '2. 点击"导入" → "自动同步"\n'
    '3. 输入Swagger文档地址：http://localhost:8080/swagger-ui/index.html\n'
    '4. 或输入OpenAPI JSON地址：http://localhost:8080/v3/api-docs\n'
    '5. 点击"立即导入"'
)
doc.add_paragraph('方式二：导入YAML/JSON文件')
add_code_block(
    '1. 点击"导入" → "导入OpenAPI/Swagger"\n'
    '2. 选择上一步生成的YAML文件\n'
    '3. 确认导入即可'
)

add_heading('3.2.2 接口分类管理', level=3)
doc.add_paragraph('推荐对接口按模块分类管理：')
categories = [
    ['认证管理', '/login', '登录、登出、获取用户信息'],
    ['用户管理', '/system/user/**', '用户CRUD操作'],
    ['入住管理', '/nursing/checkIn/**', '入住登记、查询、退住'],
    ['护理管理', '/nursing/**', '护理等级、计划、项目'],
    ['床位管理', '/nursing/{floor,room,bed}/**', '楼层、房间、床位管理'],
    ['监测管理', '/nursing/monitor/**', '设备告警、监测数据'],
    ['系统管理', '/system/**', '角色、菜单、字典等'],
    ['测试管理', '/test/case/**', '测试用例、缺陷管理'],
]
add_table(['分类', '接口路径', '说明'], categories)

add_heading('3.3 接口调试与用例管理', level=2)

add_heading('3.3.1 环境管理', level=3)
doc.add_paragraph('配置不同的测试环境，方便切换：')
add_code_block(
    '环境配置：\n'
    '├── 开发环境: http://localhost:8080\n'
    '├── 测试环境: http://test-api.zzyl.com\n'
    '├── 生产环境: https://api.zzyl.com\n\n'
    '全局变量设置：\n'
    '- base_url: http://localhost:8080\n'
    '- token: {{login_token}}  # 从登录接口自动提取\n'
    '- admin_user: admin\n'
    '- admin_pass: admin123'
)

add_heading('3.3.2 登录接口调试', level=3)
add_code_block(
    '接口：POST /login\n'
    '请求头：Content-Type: application/json\n\n'
    '请求体：\n'
    '{\n'
    '  "username": "admin",\n'
    '  "password": "admin123",\n'
    '  "code": "1",\n'
    '  "uuid": "{{captcha_uuid}}"\n'
    '}\n\n'
    '后置操作（提取token）：\n'
    '// 将token保存为环境变量\n'
    'pm.environment.set("token", pm.response.json().token)\n\n'
    '后置操作（断言）：\n'
    'pm.test("登录成功", function() {\n'
    '    pm.response.to.have.status(200);\n'
    '    pm.expect(pm.response.json().code).to.eql(200);\n'
    '    pm.exist(pm.response.json().token);\n'
    '});'
)

add_heading('3.3.3 设置全局Token', level=3)
doc.add_paragraph('在Apifox中设置全局前置/后置脚本，自动获取Token：')
add_code_block(
    '// 全局前置脚本 - 检查并刷新Token\n'
    'const token = pm.environment.get("token");\n'
    'const expired = pm.environment.get("token_expired");\n\n'
    'if (!token || Date.now() > Number(expired)) {\n'
    '    // 发送登录请求获取新Token\n'
    '    pm.sendRequest({\n'
    '        url: pm.environment.get("base_url") + "/login",\n'
    '        method: "POST",\n'
    '        header: {"Content-Type": "application/json"},\n'
    '        body: {\n'
    '            mode: "raw",\n'
    '            raw: JSON.stringify({\n'
    '                username: "admin",\n'
    '                password: "admin123",\n'
    '                code: "1",\n'
    '                uuid: "test"\n'
    '            })\n'
    '        }\n'
    '    }, function(err, res) {\n'
    '        if (res.json().code === 200) {\n'
    '            pm.environment.set("token", res.json().token);\n'
    '            pm.environment.set("token_expired", Date.now() + 3600000);\n'
    '        }\n'
    '    });\n'
    '}'
)

add_heading('3.4 自动化测试', level=2)

add_heading('3.4.1 创建测试场景', level=3)
doc.add_paragraph('场景测试可模拟真实业务流程，将多个接口串联执行：')
add_code_block(
    '场景1: 用户管理完整流程\n'
    '├── 步骤1: 登录 (POST /login)\n'
    '├── 步骤2: 获取用户列表 (GET /system/user/list)\n'
    '├── 步骤3: 新增用户 (POST /system/user)\n'
    '│   └── 提取userId\n'
    '├── 步骤4: 查询用户详情 (GET /system/user/{userId})\n'
    '├── 步骤5: 修改用户 (PUT /system/user)\n'
    '└── 步骤6: 删除用户 (DELETE /system/user/{userId})\n\n'
    '场景2: 入住管理完整流程\n'
    '├── 步骤1: 登录\n'
    '├── 步骤2: 获取楼层列表 (GET /nursing/floor/list)\n'
    '├── 步骤3: 获取房间列表 (GET /nursing/room/list)\n'
    '├── 步骤4: 新增入住登记 (POST /nursing/checkIn)\n'
    '├── 步骤5: 查询入住详情 (GET /nursing/checkIn/{id})\n'
    '└── 步骤6: 退住办理 (POST /nursing/checkOut)'
)

add_heading('3.4.2 断言编写', level=3)
add_code_block(
    '// 基础断言\n'
    'pm.test("状态码为200", () => {\n'
    '    pm.response.to.have.status(200);\n'
    '});\n\n'
    '// JSON结构验证\n'
    'pm.test("返回数据格式正确", () => {\n'
    '    const json = pm.response.json();\n'
    '    pm.expect(json.code).to.eql(200);\n'
    '    pm.expect(json).to.have.property("rows");\n'
    '});\n\n'
    '// 响应时间验证（性能测试）\n'
    'pm.test("响应时间小于1000ms", () => {\n'
    '    pm.expect(pm.response.responseTime).to.be.below(1000);\n'
    '});\n\n'
    '// 字段值验证\n'
    'pm.test("用户名验证", () => {\n'
    '    const json = pm.response.json();\n'
    '    pm.expect(json.user.userName).to.eql("admin");\n'
    '});'
)

add_heading('3.4.3 数据驱动测试', level=3)
doc.add_paragraph('使用CSV或JSON数据文件实现参数化测试，覆盖多种测试数据。')

add_heading('3.5 项目实战：养老系统接口测试', level=2)

add_heading('3.5.1 认证接口测试', level=3)
test_cases = [
    ['API-TC-001', '正确用户名密码登录', '/login', '200', '验证成功获取token'],
    ['API-TC-002', '错误密码登录', '/login', '401', '验证返回错误码和提示'],
    ['API-TC-003', '空用户名登录', '/login', '400', '验证参数校验'],
    ['API-TC-004', '获取用户信息', '/getInfo', '200', '验证返回用户详细信息'],
    ['API-TC-005', '退出登录', '/logout', '200', '验证登录状态注销'],
]
add_table(['用例编号', '用例名称', '接口', '期望状态码', '验证点'], test_cases)

add_heading('3.5.2 业务接口测试', level=3)
biz_cases = [
    ['API-TC-101', '查询用户列表', '/system/user/list', '分页参数验证、数据格式'],
    ['API-TC-102', '新增用户', '/system/user', '必填项校验、字段长度验证'],
    ['API-TC-103', '入住登记', '/nursing/checkIn', '老人信息完整、关联床位'],
    ['API-TC-104', '护理等级列表', '/nursing/level/list', '数据完整性、排序正确'],
    ['API-TC-105', '告警记录列表', '/nursing/monitor/alert/list', '时间范围过滤、状态过滤'],
]
add_table(['用例编号', '用例名称', '接口', '验证点'], biz_cases)

doc.add_page_break()

# =============================================
# 第四章 禅道缺陷管理
# =============================================
add_heading('第四章  禅道缺陷管理', level=1)

add_heading('4.1 禅道简介', level=2)
doc.add_paragraph(
    '禅道（ZenTao）是一款国产开源的项目管理软件，覆盖产品管理、项目管理、质量管理、'
    '文档管理等功能。在测试领域，禅道主要用于缺陷管理和测试用例管理。'
)

doc.add_paragraph('禅道核心概念：')
zt_concepts = [
    ['产品', '项目的产品维度，如"慧享养老服务平台"'],
    ['项目', '具体的开发项目，如"v1.0 入住管理模块开发"'],
    ['需求', '产品功能需求，关联测试用例'],
    ['用例', '测试用例，关联需求和缺陷'],
    ['缺陷(Bug)', '测试过程中发现的问题'],
    ['版本', '产品/项目的版本标识'],
    ['构建', '可测试的软件构建版本'],
]
add_table(['概念', '说明'], zt_concepts)

add_heading('4.2 缺陷生命周期', level=2)
doc.add_paragraph('一个完整的缺陷生命周期如下：')

add_code_block(
    '缺陷生命周期流转图：\n\n'
    '                   ┌─────────────┐\n'
    '                   │   新建(New)  │\n'
    '                   └──────┬──────┘\n'
    '                          │\n'
    '                          ▼\n'
    '                   ┌─────────────┐\n'
    '          ┌────────┤ 已确认(Confirmed)├────────┐\n'
    '          │        └─────────────┘        │\n'
    '          │                              │\n'
    '          ▼                              ▼\n'
    '  ┌──────────────┐              ┌──────────────┐\n'
    '  │ 已解决(Resolved)│              │ 不予解决(Will Not Fix)│\n'
    '  └──────┬───────┘              └──────┬───────┘\n'
    '         │                              │\n'
    '         ▼                              │\n'
    '  ┌──────────────┐                      │\n'
    '  │ 已关闭(Closed) │◄─────────────────────┘\n'
    '  └──────────────┘\n'
    '         │\n'
    '         ▼ (未通过验证)\n'
    '  ┌──────────────┐\n'
    '  │ 激活(Activated)├───→ 重新解决\n'
    '  └──────────────┘'
)

add_heading('4.3 缺陷提交规范', level=2)

add_heading('4.3.1 缺陷填写规范', level=3)
doc.add_paragraph('提交高质量缺陷报告是测试工程师的核心能力。以下是缺陷报告的必填字段和建议：')

fields = [
    ['标题', '简明扼要描述问题', '【登录】使用正确密码登录失败，提示"系统异常"'],
    ['所属产品', '选择对应的产品', '慧享养老服务平台'],
    ['所属模块', '选择对应的功能模块', '系统管理-登录'],
    ['所属项目', '选择对应的项目', 'v1.0'],
    ['影响版本', '发现缺陷的版本', 'v1.0.0'],
    ['当前指派', '指派给对应的开发人员', '张三'],
    ['缺陷类型', '代码错误/设计缺陷/需求问题/兼容性', '代码错误'],
    ['严重程度', '致命/严重/一般/建议', '严重'],
    ['优先级', 'P0/P1/P2/P3', 'P1'],
    ['重现步骤', '详细描述复现步骤', '1. 打开登录页 2. 输入正确账号密码 3. 点击登录'],
    ['实际结果', '描述实际发生的现象', '提示"系统异常"，无法登录'],
    ['期望结果', '描述应该期望的结果', '登录成功，跳转到首页'],
    ['附件', '截图/日志/录屏', 'error_log.txt, screenshot.png'],
]
add_table(['字段', '说明', '示例'], fields)

add_heading('4.3.2 严重程度分级', level=3)
severity = [
    ['P0-致命', '系统崩溃/数据丢失/主流程不可用', '登录失败导致无法使用系统'],
    ['P1-严重', '主要功能失效，但有替代方案', '入住登记保存失败'],
    ['P2-一般', '功能异常但不影响核心流程', '搜索结果显示不完整'],
    ['P3-建议', '界面优化/功能建议/体验改进', '按钮颜色不美观、文案建议'],
]
add_table(['等级', '定义', '示例'], severity)

add_heading('4.3.3 缺陷提交示例', level=3)
add_code_block(
    '缺陷标题：【床位管理】点击删除按钮后未弹出确认提示直接删除\n\n'
    '严重程度：一般 (P2)\n'
    '优先级：P2\n'
    '模块：床位管理\n\n'
    '重现步骤：\n'
    '1. 登录慧享养老服务平台\n'
    '2. 进入"床位管理"模块\n'
    '3. 选择任意一个空闲床位\n'
    '4. 点击"删除"按钮\n\n'
    '实际结果：点击删除按钮后，床位被直接删除，没有任何确认提示\n\n'
    '期望结果：点击删除按钮后，应弹出确认对话框，用户确认后才执行删除操作\n\n'
    '附件：delete_bug_screenshot.png'
)

add_heading('4.4 项目实战：养老系统缺陷管理', level=2)

add_heading('4.4.1 缺陷跟踪流程', level=3)
doc.add_paragraph('在慧享养老服务平台项目中，缺陷管理流程如下：')
add_code_block(
    '1. 测试执行阶段发现缺陷\n'
    '   ↓\n'
    '2. 在禅道中提交缺陷(Bug)，填写完整信息\n'
    '   ↓\n'
    '3. 开发人员确认缺陷 (分配 → 确认)\n'
    '   ↓\n'
    '4. 开发修复缺陷 (解决 → 填写解决方案)\n'
    '   ↓\n'
    '5. 测试人员回归验证 (验证通过 → 关闭)\n'
    '   ↓\n'
    '6. 回归不通过 → 激活(Activate)，返回步骤3'
)

add_heading('4.4.2 缺陷统计与分析', level=3)
add_code_block(
    '# 缺陷统计SQL示例（禅道数据库）\n'
    'SELECT\n'
    '    module AS "模块",\n'
    '    COUNT(*) AS "缺陷数",\n'
    '    SUM(CASE WHEN severity = 1 THEN 1 ELSE 0 END) AS "致命",\n'
    '    SUM(CASE WHEN severity = 2 THEN 1 ELSE 0 END) AS "严重",\n'
    '    SUM(CASE WHEN severity = 3 THEN 1 ELSE 0 END) AS "一般",\n'
    '    SUM(CASE WHEN severity = 4 THEN 1 ELSE 0 END) AS "建议"\n'
    'FROM zt_bug\n'
    'WHERE project = 1\n'
    'GROUP BY module\n'
    'ORDER BY 缺陷数 DESC;'
)

doc.add_paragraph()
doc.add_paragraph('禅道中常用的缺陷分析维度：')
analysis = [
    ['按模块统计', '识别高缺陷模块，评估代码质量'],
    ['按严重程度', '评估系统稳定性'],
    ['按引入阶段', '改进研发流程'],
    ['按处理人', '跟踪开发人员修复效率'],
    ['按版本统计', '对比各版本质量趋势'],
    ['缺陷密度', '每千行代码缺陷数（KLOCC）'],
]
add_table(['分析维度', '说明'], analysis)

doc.add_page_break()

# =============================================
# 第五章 测试全流程实战
# =============================================
add_heading('第五章  测试全流程实战', level=1)

doc.add_paragraph(
    '本章以慧享养老服务平台的"入住管理模块"为例，完整演示从测试计划到测试报告的全流程。'
)

add_heading('5.1 测试计划', level=2)
doc.add_paragraph('测试计划是测试工作的蓝图，包含以下要素：')

add_code_block(
    '测试计划大纲\n'
    '========================================\n'
    '1. 测试范围\n'
    '   - 功能测试：入住登记、退住办理、信息查询\n'
    '   - 接口测试：CheckInController 所有接口\n'
    '   - UI测试：表单验证、页面跳转、数据展示\n\n'
    '2. 测试策略\n'
    '   - 功能测试：等价类划分 + 边界值分析\n'
    '   - 接口测试：参数校验 + 业务场景 + 异常测试\n'
    '   - UI自动化：Selenium + pytest (覆盖核心流程)\n\n'
    '3. 测试环境\n'
    '   - Web：Chrome 120+, Windows 11\n'
    '   - 后端：JDK 17, Spring Boot 3.5.x\n'
    '   - 数据库：MySQL 8.0\n'
    '   - 缓存：Redis 7.x\n\n'
    '4. 测试数据\n'
    '   - 正常数据：5组合法入住信息\n'
    '   - 异常数据：空字段、超长字段、重复数据\n'
    '   - 边界数据：年龄边界、金额边界\n\n'
    '5. 测试进度\n'
    '   - 第1天：测试环境搭建 + 用例设计\n'
    '   - 第2天：功能测试执行\n'
    '   - 第3天：接口测试 + 自动化脚本\n'
    '   - 第4天：回归测试 + 缺陷跟踪\n'
    '   - 第5天：测试报告编写\n\n'
    '6. 准出标准\n'
    '   - 所有P0/P1用例通过\n'
    '   - 无致命/严重缺陷\n'
    '   - 核心业务流程跑通\n'
    '   - 测试报告评审通过'
)

add_heading('5.2 测试用例设计', level=2)

add_heading('5.2.1 等价类划分法', level=3)
doc.add_paragraph('以入住登记功能的"年龄"字段为例：')
add_code_block(
    '年龄字段等价类划分：\n'
    '┌───────────────┬─────────────────┬──────────────┐\n'
    '│ 有效等价类     │ 无效等价类       │ 边界值       │\n'
    '├───────────────┼─────────────────┼──────────────┤\n'
    '│ 18~150        │ <18            │ 18, 150      │\n'
    '│               │ >150            │ 17, 151      │\n'
    '│               │ 非数字          │ abc, 12.5    │\n'
    '│               │ 为空            │ null, ""     │\n'
    '└───────────────┴─────────────────┴──────────────┘'
)

add_heading('5.2.2 边界值分析法', level=3)
add_code_block(
    '入住登记 - 字段边界值测试用例：\n\n'
    '姓名：1字、20字（正常）| 0字、21字（异常）\n'
    '年龄：18、150（正常）| 17、151（异常）\n'
    '电话：11位数字（正常）| 10位、12位、含字母（异常）\n'
    '身份证：18位（正常）| 17位、19位、含非法字符（异常）\n'
    '金额(押金)：0.01（最小正数）| 0、负数（异常）'
)

add_heading('5.2.3 场景法', level=3)
add_code_block(
    '入住管理场景：\n\n'
    '基本流（Happy Path）：\n'
    '  登录系统 → 进入入住管理 → 点击新增\n'
    '  → 填写老人信息 → 选择床位 → 保存 → 查看列表\n\n'
    '备选流1：\n'
    '  填写信息 → 未选择床位直接保存 → 提示"请选择床位"\n\n'
    '备选流2：\n'
    '  填写信息 → 选择已被占用的床位 → 提示"床位已被占用"\n\n'
    '备选流3：\n'
    '  查看列表 → 编辑入住信息 → 修改关键字段 → 保存\n\n'
    '异常流：\n'
    '  提交表单 → 网络中断 → 提示"网络异常" → 恢复网络 → 重试'
)

add_heading('5.3 测试执行', level=2)

add_heading('5.3.1 功能测试执行', level=3)
doc.add_paragraph('执行测试用例，记录结果：')
execution = [
    ['TC-001', '入住登记-成功', 'P0', '通过', '填写完整信息，保存成功'],
    ['TC-002', '入住登记-姓名为空', 'P1', '通过', '提示"姓名不能为空"'],
    ['TC-003', '入住登记-无效年龄', 'P1', '失败', '输入-5未拦截（需提Bug）'],
    ['TC-004', '入住登记-重复身份证', 'P2', '通过', '提示"身份证已存在"'],
    ['TC-005', '入住登记-特殊字符', 'P2', '阻塞', '输入<scirpt>标签导致页面异常'],
]
add_table(['用例编号', '用例名称', '优先级', '执行结果', '备注'], execution)

add_heading('5.3.2 接口测试执行', level=3)
doc.add_paragraph('使用Apifox执行接口测试并记录结果：')
api_exec = [
    ['POST /login', '200', '通过', '0.3s', '获取token成功'],
    ['GET /user/list', '200', '通过', '0.5s', '数据格式正确'],
    ['POST /checkIn', '201', '通过', '0.8s', '入住记录创建成功'],
    ['POST /checkIn', '400', '通过', '0.2s', '必填项校验生效'],
    ['GET /checkIn/{id}', '200', '通过', '0.3s', '返回详情数据完整'],
    ['DELETE /checkIn/{id}', '200', '通过', '0.4s', '删除成功'],
]
add_table(['接口', '状态码', '结果', '响应时间', '说明'], api_exec)

add_heading('5.3.3 自动化测试执行', level=3)
doc.add_paragraph('使用 pytest 执行自动化测试脚本：')
add_code_block(
    '# 终端运行\n'
    'cd zzyl-autotest\n'
    'pytest testcases/ -v --html=reports/report.html\n\n'
    '# 执行结果示例\n'
    '============================= test session starts =============================\n'
    'collecting ... collected 8 items\n\n'
    'test_login.py::TestLogin::test_login_success PASSED                    [12%]\n'
    'test_login.py::TestLogin::test_login_wrong_password PASSED             [25%]\n'
    'test_login.py::TestLogin::test_login_empty_fields PASSED              [37%]\n'
    'test_checkin.py::TestCheckin::test_checkin_list PASSED                [50%]\n'
    'test_checkin.py::TestCheckin::test_checkin_search PASSED              [62%]\n'
    'test_checkin.py::TestCheckin::test_checkin_form FAILED                [75%]\n'
    'test_api.py::TestAPI::test_login_api PASSED                           [87%]\n'
    'test_api.py::TestAPI::test_get_user_info PASSED                       [100%]\n\n'
    '======================= 7 passed, 1 failed in 45.2s ========================'
)

add_heading('5.4 缺陷跟踪', level=2)
doc.add_paragraph('测试执行过程中发现的缺陷需提交到禅道进行跟踪：')
bugs = [
    ['BUG-001', '年龄字段未校验负值', 'P1', '严重', '已修复', '开发修复完成'],
    ['BUG-002', 'XSS注入未过滤', 'P1', '严重', '修复中', '正在修复'],
    ['BUG-003', '列表翻页后数据重复', 'P2', '一般', '新建', '未开始处理'],
    ['BUG-004', '退住保存后页面未刷新', 'P2', '一般', '已关闭', '验证通过'],
    ['BUG-005', '手机号格式校验不完整', 'P3', '建议', '不予解决', '后续版本优化'],
]
add_table(['Bug编号', '标题', '优先级', '严重程度', '状态', '备注'], bugs)

add_heading('5.5 测试报告', level=2)
doc.add_paragraph('测试报告是测试工作的总结和产出物，包含：')

add_heading('5.5.1 报告结构', level=3)
add_code_block(
    '慧享养老服务平台 v1.0 - 入住管理模块测试报告\n'
    '============================================\n\n'
    '1. 测试概述\n'
    '   测试时间：2024年6月1日 - 6月5日\n'
    '   测试范围：入住管理模块（全部功能）\n'
    '   测试环境：Windows 11 / Chrome 120\n\n'
    '2. 测试统计\n'
    '   ├── 用例总数：50\n'
    '   ├── 通过：42 (84%)\n'
    '   ├── 失败：5 (10%)\n'
    '   ├── 阻塞：3 (6%)\n'
    '   └── 通过率：84%\n\n'
    '3. 缺陷统计\n'
    '   ├── 提交缺陷：8\n'
    '   ├── 已修复：5\n'
    '   ├── 已验证关闭：4\n'
    '   └── 遗留问题：3\n\n'
    '4. 测试结论\n'
    '   └── 建议：修复所有严重缺陷后进行第二轮回归测试\n'
    '       达到"所有P0用例通过、无致命缺陷"标准后准出'
)

add_heading('5.5.2 缺陷分布分析', level=3)
add_code_block(
    '按模块缺陷分布：\n'
    '├── 入住登记：3个 (37.5%)  ← 重点关注\n'
    '├── 退住办理：2个 (25.0%)\n'
    '├── 列表查询：2个 (25.0%)\n'
    '└── 数据导出：1个 (12.5%)\n\n'
    '按严重程度分布：\n'
    '├── 致命：0 (0%)\n'
    '├── 严重：3 (37.5%)\n'
    '├── 一般：4 (50.0%)\n'
    '└── 建议：1 (12.5%)'
)

doc.add_page_break()

# =============================================
# 附录：常见面试题
# =============================================
add_heading('附录：常见测试面试题及答案', level=1)

interview_qa = [
    ('1. 自我介绍',
     '面试官你好，我叫xxx，今年24岁。我专注于软件测试领域，熟练掌握功能测试、接口测试、自动化测试。'
     '我熟悉Python语言，能使用Selenium进行Web自动化测试，使用Apifox进行接口测试，使用禅道进行缺陷管理。'
     '在之前的项目中，我负责慧享养老服务平台的测试工作，参与了需求评审、用例设计、测试执行、缺陷跟踪等全流程。'
     '我逻辑思维强，细心严谨，善于发现Bug，能够和产品、开发高效沟通，推动问题解决。'),

    ('2. 请介绍一下你的项目',
     '我最近参与的是一个养老管理系统项目，我主要负责后台管理系统的测试工作。'
     '后台管理系统供养老院员工使用，主要模块包括：入住管理、退住管理、护理管理、床位管理、智能监测等。'
     '项目基于若依框架（RuoYi-Vue）开发，后端使用Spring Boot、MyBatis、MySQL、Redis，前端使用Vue3 + Element Plus。'
     '我在项目中负责全流程测试工作，包括测试计划编写、测试用例设计、功能测试、接口测试、自动化测试、缺陷跟踪和测试报告编写。'),

    ('3. 测试流程是什么？',
     '标准测试流程包括以下阶段：\n'
     '1）需求分析：参与需求评审，理解业务需求\n'
     '2）测试计划：制定测试计划，明确范围、策略、资源\n'
     '3）用例设计：使用等价类、边界值、场景法等设计测试用例\n'
     '4）测试执行：执行功能测试、接口测试、自动化测试\n'
     '5）缺陷管理：在禅道中提交并跟踪缺陷\n'
     '6）回归测试：修复后进行回归验证\n'
     '7）测试报告：汇总测试结果，输出测试报告\n'
     '8）上线跟进：上线后跟进线上问题'),

    ('4. 如何设计测试用例？',
     '我会结合多种方法设计测试用例：\n'
     '1）等价类划分法：将输入数据划分为有效和无效等价类\n'
     '2）边界值分析法：重点关注边界值（如年龄18~150，测试17、18、150、151）\n'
     '3）场景法：覆盖基本流、备选流、异常流（如入住管理完整流程）\n'
     '4）错误推测法：根据经验推测可能出错的地方（如特殊字符、XSS注入）\n'
     '5）正交实验法：多参数组合时使用正交表（如搜索条件的多条件组合）\n'
     '总之先覆盖核心业务流程（P0用例），再覆盖异常场景和边界条件。'),

    ('5. 如何定位Bug？',
     '定位Bug我遵循以下步骤：\n'
     '1）确认Bug环境：确定是哪个环境、哪个版本\n'
     '2）复现Bug：按照步骤尝试稳定复现\n'
     '3）区分前后端：打开浏览器开发者工具(F12)，查看Network请求\n'
     '   - 前端问题：请求未发出、请求参数错误、页面渲染异常\n'
     '   - 后端问题：接口返回数据错误、状态码异常、响应时间过长\n'
     '4）查看日志：请求后端查看日志（tail -f, ELK）\n'
     '5）缩小范围：通过二分法缩小问题范围\n'
     '6）记录信息：截图、保存日志、记录操作步骤'),

    ('6. Selenium中如何等待元素加载？',
     '我主要使用三种等待方式：\n'
     '1）隐式等待(driver.implicitly_wait)：全局等待，每次查找元素时最多等待N秒\n'
     '2）显式等待(WebDriverWait + expected_conditions)：针对特定元素的等待\n'
     '   如：WebDriverWait(driver, 10).until(EC.presence_of_element_located((By.ID, "xxx")))\n'
     '3）强制等待(time.sleep)：不推荐，仅在调试时使用\n'
     '在工作中我主要使用显式等待，它效率高、针对性强。'),

    ('7. 如何做接口测试？',
     '我使用Apifox进行接口测试，主要分以下步骤：\n'
     '1）接口文档：从后端获取Swagger/OpenAPI文档导入Apifox\n'
     '2）环境配置：配置开发/测试/生产环境\n'
     '3）接口调试：逐个调试接口，验证请求参数和响应数据\n'
     '4）用例编写：为每个接口编写测试用例，包含正常和异常场景\n'
     '5）断言设置：验证状态码、响应时间、返回数据格式和内容\n'
     '6）场景测试：将多个接口串联成业务流程\n'
     '7）自动化测试：使用Apifox的自动化测试功能定时执行\n'
     '8）CI集成：集成到Jenkins实现持续测试'),

    ('8. 如何使用Selenium处理弹窗？',
     'Selenium处理弹窗的方法：\n'
     '1）Alert弹窗：\n'
     '   alert = driver.switch_to.alert\n'
     '   alert.accept()  # 确定\n'
     '   alert.dismiss()  # 取消\n'
     '   alert.send_keys("text")  # 输入文本\n'
     '2）确认框(confirm)：同上\n'
     '3）提示框(prompt)：同上\n'
     '4）自定义弹窗（Element Plus Dialog）：通过XPath定位弹窗元素\n'
     '   driver.find_element(By.XPATH, "//div[contains(@class, \'el-dialog\')]//button[contains(text(), \'确定\')]")'),

    ('9. 如何处理上传文件？',
     '文件上传的自动化处理方式：\n'
     '1）input标签：直接使用send_keys上传\n'
     '   driver.find_element(By.XPATH, "//input[@type=\'file\']").send_keys("C:\\\\path\\\\file.pdf")\n'
     '2）非input标签（如自定义上传组件）：\n'
     '   - 使用pyautogui模拟键盘操作\n'
     '   - 或使用AutoIt工具\n'
     '3）对于慧享养老系统的体检报告上传：上传PDF后调用AI分析，自动化时需等待AI处理完成'),

    ('10. 禅道中缺陷的状态有哪些？',
     '禅道中缺陷的常见状态：\n'
     '1）新建(New)：测试人员提交缺陷\n'
     '2）已确认(Confirmed)：开发确认是Bug\n'
     '3）已解决(Resolved)：开发修复完成\n'
     '4）已关闭(Closed)：测试验证通过\n'
     '5）激活(Activated)：回归不通过，重新激活\n'
     '6）不予解决(Will Not Fix)：设计如此或非问题\n'
     '7）延期处理(Postponed)：当前版本不修，后续版本处理'),

    ('11. 什么是PO模式？有什么优点？',
     'PO（Page Object）模式是Web自动化的设计模式，将每个页面封装成一个类：\n'
     '- 元素定位放在类中\n'
     '- 页面操作封装成方法\n\n'
     '优点：\n'
     '1）代码复用：页面元素变化时只需修改一处\n'
     '2）维护性：减少重复代码，提高可维护性\n'
     '3）可读性：测试用例清晰易读，关注业务逻辑\n'
     '4）分层设计：页面对象层 + 测试用例层 + 数据层'),

    ('12. 如何保证测试覆盖率？',
     '我通过以下方法保证测试覆盖率：\n'
     '1）需求覆盖率：确保每个需求点都有对应的测试用例\n'
     '2）代码覆盖率：使用JaCoCo工具统计行覆盖、分支覆盖\n'
     '3）接口覆盖率：确保所有API接口都被覆盖测试\n'
     '4）业务场景覆盖率：覆盖所有业务流程（基本流+备选流）\n'
     '5）自动化覆盖率：核心业务实现自动化覆盖\n'
     '6）交叉评审：测试用例由团队交叉评审，查漏补缺'),

    ('13. 性能测试关注哪些指标？',
     '性能测试主要关注：\n'
     '1）响应时间：接口/页面的响应时长（目标<2s）\n'
     '2）吞吐量(TPS/QPS)：每秒处理的事务数/请求数\n'
     '3）并发用户数：系统能同时承载的用户数\n'
     '4）错误率：请求失败的比例（目标<0.1%）\n'
     '5）CPU使用率：<70%\n'
     '6）内存使用率：<80%\n'
     '7）数据库连接数、慢查询'),

    ('14. MySQL如何优化慢查询？',
     '优化慢查询的步骤：\n'
     '1）开启慢查询日志，定位慢SQL\n'
     '2）使用EXPLAIN分析执行计划\n'
     '3）检查索引使用情况（key、key_len、type）\n'
     '4）优化手段：\n'
     '   - 添加合适索引（关注联合索引最左匹配）\n'
     '   - 避免SELECT *，只查询需要的字段\n'
     '   - 避免在WHERE中使用函数或隐式类型转换\n'
     '   - 拆分大查询为小查询\n'
     '   - 使用分页优化（limit深分页问题）\n'
     '5）验证优化效果'),

    ('15. 如何处理翻页中的"重复数据"缺陷？',
     '翻页重复数据是常见缺陷，排查步骤：\n'
     '1）确认排序字段是否唯一（如只有create_time排序，同时间创建的数据会乱序）\n'
     '2）建议排序加上唯一字段（如ORDER BY create_time DESC, id DESC）\n'
     '3）检查分页SQL中OFFSET计算是否正确\n'
     '4）检查Redis缓存与数据库数据一致性\n'
     '5）提交Bug时附上翻页前后的截图和SQL日志'),
]

for q, a in interview_qa:
    add_heading(q, level=2)
    doc.add_paragraph(a)
    doc.add_paragraph()

# =============================================
# 保存文档
# =============================================
output_path = "C:/Users/ASUS/Desktop/生产实习/zyyl/慧享养老服务平台-测试全流程操作手册.docx"
doc.save(output_path)
print(f"测试操作手册已生成: {output_path}")
