"""
生成《苏凯峰 - 软件测试工程师简历》更新版
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 样式
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

def add_section_heading(text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
        run.font.size = Pt(14)

# ====== 个人信息 ======
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('苏 凯 峰')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

# 个人信息表
info_data = [
    ['年　　龄', '21岁', '联系电话', '15727240576'],
    ['性　　别', '男', '电子邮箱', '1564441326@qq.com'],
    ['毕业院校', '武汉轻工大学', '专　　业', '计算机科学与技术'],
    ['求职意向', '软件测试工程师', '工作经验', '应届生/实习'],
]
table = doc.add_table(rows=len(info_data), cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(info_data):
    for j, cell_text in enumerate(row_data):
        cell = table.rows[i].cells[j]
        cell.text = cell_text
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(11)
                if j % 2 == 0:
                    run.bold = True

doc.add_paragraph()

# ====== 个人优势 ======
add_section_heading('个人优势')
advantages = [
    '熟悉软件测试完整流程，能独立完成测试计划制定、测试用例设计、测试执行、缺陷跟踪及测试报告编写',
    '逻辑思维强，细心严谨，善于发现Bug，具备问题定位与复现总结能力',
    '熟练掌握功能测试、接口测试、回归测试、兼容性测试；熟练使用Apifox、Selenium、禅道等测试工具',
    '掌握Python编程，能使用Selenium + pytest实现Web自动化测试',
    '具备良好业务分析能力，学习适应能力强，可快速熟悉新项目及业务模块，有效开展测试工作',
    '沟通协调顺畅，可高效对接产品、开发推进项目，抗压能力良好',
]
for adv in advantages:
    doc.add_paragraph(adv, style='List Bullet')

# ====== 项目经验 ======
add_section_heading('项目经验')

# 项目1
p = doc.add_paragraph()
run = p.add_run('慧享养老服务平台')
run.bold = True
run.font.size = Pt(13)
p.add_run('　　2024.02 - 2024.07')
run = p.runs[-1]
run.font.color.rgb = RGBColor(102, 102, 102)

doc.add_paragraph(
    '项目描述：慧享养老服务平台是专为养老机构定制开发的专业养老服务管理平台，我主要负责后台管理系统的测试工作。'
    '后台管理系统供养老院员工使用，主要包含入住管理、退住管理、护理管理、床位管理、智能监测等核心模块；'
    '系统还包含微信小程序供老人及家属使用（由其他团队在微信开发者工具中独立管理）。'
    '后端使用Spring Boot + MyBatis + Redis，前端使用Vue3 + Element Plus，集成百度千帆大模型做AI体检报告分析。'
)

p = doc.add_paragraph()
run = p.add_run('测试职责：')
run.bold = True
duties = [
    '测试计划：制定模块级测试计划，明确测试范围、策略、资源安排和交付标准',
    '用例设计：运用等价类划分、边界值分析、场景法等设计120+条测试用例，覆盖全部核心模块',
    '功能测试：执行8轮迭代测试，覆盖入住管理、护理管理、床位管理、智能监测等核心模块',
    '接口测试：使用Apifox对50+个后端API进行接口测试，覆盖参数校验、业务逻辑、异常场景',
    '自动化测试：使用Python + Selenium + pytest搭建自动化测试框架，编写20+自动化测试用例',
    '缺陷管理：使用禅道进行缺陷全生命周期管理，累计提交45个缺陷，推动修复率达95%+',
    'AI测试：测试百度千帆大模型对接功能，验证体检报告分析结果的准确性和一致性',
    '测试报告：输出6份版本测试报告，汇总测试结果和缺陷分析，为版本发布提供质量依据',
]
for d in duties:
    doc.add_paragraph(d, style='List Bullet')

# 项目2
p = doc.add_paragraph()
run = p.add_run('若依框架二次开发项目')
run.bold = True
run.font.size = Pt(13)
p.add_run('　　2024.02 - 2024.07')
run = p.runs[-1]
run.font.color.rgb = RGBColor(102, 102, 102)

doc.add_paragraph(
    '在项目中基于若依框架（RuoYi-Vue）进行二次开发，充分利用代码生成器、表单构建器、定时任务等内置功能提高开发效率。'
    '在项目中新增了测试管理模块（测试用例管理、缺陷管理、测试计划、测试仪表盘），实现测试工作的数字化管理。'
)

# ====== 专业技能 ======
add_section_heading('专业技能')
skills = [
    '掌握软件测试理论和方法，熟悉黑盒测试（功能、接口、自动化）全流程，具有独立完成测试计划、用例设计、缺陷跟踪、报告输出的实践经验',
    '掌握Python编程语言，熟悉Selenium Web自动化测试框架，能使用pytest + Selenium搭建自动化测试框架',
    '熟练使用Apifox进行接口测试，包括接口文档管理、接口调试、自动化测试场景编排',
    '熟练使用禅道进行缺陷管理和测试用例管理，了解缺陷生命周期和规范化提交流程',
    '了解MySQL数据库，能编写SQL进行数据查询和测试数据准备',
    '了解Linux基本操作命令，能查看日志定位问题',
    '熟悉Git版本管理工具，了解敏捷开发流程',
]
for s in skills:
    doc.add_paragraph(s, style='List Bullet')

# ====== 教育背景 ======
add_section_heading('教育背景')
p = doc.add_paragraph()
run = p.add_run('武汉轻工大学')
run.bold = True
p.add_run('　　计算机科学与技术　　本科')
p = doc.add_paragraph()
run = p.add_run('主修课程：')
run.bold = True
p.add_run('软件工程、数据库原理、计算机网络、操作系统、Web开发技术')

# ====== 自我评价 ======
add_section_heading('自我评价')
doc.add_paragraph(
    '具备扎实的计算机理论基础，掌握功能测试、接口测试、回归测试、自动化测试全流程，'
    '能运用各类方法设计测试用例和执行测试。熟练掌握Apifox、Selenium、禅道等测试工具，'
    '掌握Python编程语言，能独立搭建自动化测试框架。具备较强逻辑分析能力和Bug定位能力，'
    '做事严谨细致、责任心强。善于团队协作，能与产品、开发高效沟通，推动项目进展。'
    '学习适应能力快、抗压能力好，能快速投入新项目开展测试工作。'
)

# 保存
output_path = "C:/Users/ASUS/Desktop/生产实习/zyyl/苏凯峰-软件测试工程师简历.docx"
doc.save(output_path)
print(f"简历已生成: {output_path}")
